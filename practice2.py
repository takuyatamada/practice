import shutil
import docx
shutil.copy("template.txt","input.txt")
#pra2
doc = docx.Document()
doc.add_paragraph("aaa")
doc.add_paragraph("bbb")
title = "title"
doc.save(title+'.docx')